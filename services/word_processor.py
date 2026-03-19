import io
from docx import Document

class WordProcessor:
    @staticmethod
    def extract_text(docx_content: bytes) -> str:
        """
        Extracts text from a Word document (.docx) including paragraphs and tables.
        """
        doc = Document(io.BytesIO(docx_content))
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)

word_processor = WordProcessor()
